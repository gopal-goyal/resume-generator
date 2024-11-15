import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_resume(file_name):
    # Load the data from the JSON file
    with open('data.json', 'r') as f:
        data_dict = json.load(f)
    
    # Fetch the data for the given file_name
    data = data_dict.get(file_name)
    if not data:
        print(f"Error: No data found for file name: {file_name}")
        return
    
    doc = Document()
    
    # Step 1: Name and Contact Info
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]

    # Left Side
    left_cell = row.cells[0]
    left_paragraph = left_cell.paragraphs[0]
    left_paragraph.add_run(f"{data['name']}\n").bold = True
    left_paragraph.add_run(f"{data['designation']}\n{data['department']}\n{data['company']}")
    left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Right Side
    right_cell = row.cells[1]
    right_paragraph = right_cell.paragraphs[0]
    contact_info = data['contact']
    right_paragraph.add_run(f"📞 Contact Number: {contact_info['phone']}\n").bold = True
    right_paragraph.add_run(f"✉️ Email: {contact_info['email']}\n")
    right_paragraph.add_run(f"🐱 GitHub: {contact_info['github']}\n")
    right_paragraph.add_run(f"🔗 LinkedIn: {contact_info['linkedin']}\n")
    right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph()

    # Step 2: Professional Summary
    doc.add_heading('Professional Summary', level=2)
    doc.add_paragraph(data['summary'])
    doc.add_paragraph()

    # Step 3: Skills Section
    doc.add_heading('Skills', level=2)
    doc.add_paragraph('\n'.join(data['skills']))
    doc.add_paragraph()

    # Step 4: Experience Section
    doc.add_heading('Experience', level=2)
    doc.add_paragraph(data['experience'])
    doc.add_paragraph()

    # Step 5: Projects Section
    doc.add_heading('Projects', level=2)
    doc.add_paragraph('\n'.join(data['projects']))
    doc.add_paragraph()

    # Step 6: Education Section
    doc.add_heading('Education', level=2)
    table = doc.add_table(rows=1, cols=4)

    # Set column headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Degree'
    hdr_cells[1].text = 'Institute/Board'
    hdr_cells[2].text = 'CGPA/Percentage'
    hdr_cells[3].text = 'Year'

    # Add rows for each education entry
    for edu in data['education']:
        row_cells = table.add_row().cells
        row_cells[0].text = edu['degree']
        row_cells[1].text = edu['institute']
        row_cells[2].text = edu['cgpa']
        row_cells[3].text = edu['year']

    # Save the document
    doc.save(f'data/{file_name}.docx')
    print(f"Resume generated and saved as: {file_name}")

