from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create Document Object
doc = Document()

# Step 1: Name and Contact Info (Left and Right Aligned)
table = doc.add_table(rows=1, cols=2)
row = table.rows[0]

# Left Side: Name, Designation, Company Name
left_cell = row.cells[0]
left_paragraph = left_cell.paragraphs[0]
left_paragraph.add_run('Gopal Goyal\n').bold = True
left_paragraph.add_run('Software Engineer L2\nData Science Department\nGemini Solutions Pvt. Ltd.')
left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Right Side: Contact Info
right_cell = row.cells[1]
right_paragraph = right_cell.paragraphs[0]
right_paragraph.add_run('📞 Contact Number: [Your Contact Number]\n').bold = True
right_paragraph.add_run('✉️ Email: goyal11.gopal@gmail.com\n')
right_paragraph.add_run('🐱 GitHub: https://github.com/gopal-goyal\n')
right_paragraph.add_run('🔗 LinkedIn: https://www.linkedin.com/in/gopal911/\n')
right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc.add_paragraph()

# Step 2: Professional Summary
doc.add_heading('Professional Summary', level=2)
summary = doc.add_paragraph('I am a Software Engineer L2 at Gemini Solutions Pvt. Ltd., specializing in Python development, machine learning, and data science. '
                            'With a solid technical understanding and problem-solving skills, I excel in building end-to-end AI solutions that effectively address complex challenges. '
                            'My expertise in large language models (LLMs), data analysis, and software development allows me to deliver optimized codebases for production environments. '
                            'I have a proven track record of contributing to various AI and ML projects, including sports performance analytics, financial sentiment analysis, and text-to-video solutions.')

doc.add_paragraph()

# Step 3: Skills Section
doc.add_heading('Skills', level=2)
skills = '''
- Programming Languages: Python, C, C++, Dart, JavaScript, HTML, CSS
- Machine Learning: TensorFlow, PyTorch, Scikit-learn, Roboflow, Pandas, NumPy, Matplotlib, Seaborn
- Prompt Engineering: Large Language Models (LLMs), Text-to-Video, Stable Diffusion, FinBERT
- App Development: React.js, Flutter, Firebase
- Financial Data Handling: Forex sentiment analysis, Graphical comparisons, Dashboards
- Tools & Technologies: AWS (ECS), Git, Docker, Boto3, MoviePy, VSCode, Linux, Windows
'''
doc.add_paragraph(skills)

doc.add_paragraph()

# Step 4: Experience Section
doc.add_heading('Experience', level=2)
doc.add_paragraph('''As a Software Developer L2 in the Data Science Department, I contribute to various AI and ML projects, utilizing my skills in Python and data analysis to create solutions that meet client needs. 
I specialize in developing and deploying software that enhances operational efficiency and drives business growth.''')

doc.add_paragraph()

# Step 5: Projects Section
doc.add_heading('Projects', level=2)
projects = '''
- Text-to-Video Solution: Developed an end-to-end AI-powered text-to-video conversion solution using large language models and Stable Diffusion, successfully hosting the entire pipeline to deliver a seamless user experience. This project highlights my ability to integrate multiple technologies to produce impactful solutions.
- Forex Sentiment Analysis: Designed and implemented a comprehensive dashboard that effectively analyzes financial data and sentiment scoring. My work demonstrates strong capabilities in handling financial datasets and creating insightful visualizations that facilitate informed decision-making for users.
- Sports Performance Analytics: Led the training of various machine learning models, including object detection and classification, using tools like MediaPipe and SVM. I also manage the codebase for the machine learning aspect of the project, showcasing my technical expertise and ability to contribute to complex AI initiatives.
- Personal Projects: Proficient in both React and Flutter, I have developed several mobile applications, including weather apps and hidden calculator vault apps. This versatility underscores my ability to adapt and excel across different platforms and technologies.
'''
doc.add_paragraph(projects)

doc.add_paragraph()

# Step 6: Education Section (Table)
doc.add_heading('Education', level=2)
table = doc.add_table(rows=3, cols=4)

# Set column headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Degree'
hdr_cells[1].text = 'Institute/Board'
hdr_cells[2].text = 'CGPA/Percentage'
hdr_cells[3].text = 'Year'

# Row 1: B.Tech
row_cells = table.rows[1].cells
row_cells[0].text = 'B.Tech in Computer Science and Engineering'
row_cells[1].text = 'NIT Hamirpur'
row_cells[2].text = '8.38/10'
row_cells[3].text = '2019-2023'

# Row 2: XII CBSE
row_cells = table.rows[2].cells
row_cells[0].text = 'XII CBSE'
row_cells[1].text = 'Lav Kush Convent School, SGNR'
row_cells[2].text = '92.4%'
row_cells[3].text = '2019'

# Row 3: X CBSE
table.add_row()
row_cells = table.rows[3].cells
row_cells[0].text = 'X CBSE'
row_cells[1].text = 'Nosegay Public School, SGNR'
row_cells[2].text = '95%'
row_cells[3].text = '2017'

# Save the document
doc.save('Gopal_Goyal_Developer_Resume.docx')
